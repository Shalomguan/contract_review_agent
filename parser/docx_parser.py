"""
Word文档解析器
使用python-docx提取文本
"""
from docx import Document
from pathlib import Path
from typing import Optional
from .base import BaseParser, ParsedDocument


class DocxParser(BaseParser):
    """Word文档(.docx)解析器"""

    def __init__(self):
        self.supported_extensions = ['.docx']

    def parse(self, file_path: str) -> ParsedDocument:
        """解析Word文档"""
        text = self.extract_text(file_path)
        metadata = self._extract_metadata(file_path)

        return ParsedDocument(
            text=text,
            title=metadata.get('title'),
            parties=metadata.get('parties'),
            file_type='docx',
            raw_metadata=metadata
        )

    def extract_text(self, file_path: str) -> str:
        """从Word文档提取文本"""
        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return '\n\n'.join(paragraphs)

    def _extract_metadata(self, file_path: str) -> dict:
        """提取Word文档元数据"""
        doc = Document(file_path)
        metadata = {}

        # 提取核心属性
        core_props = doc.core_properties
        metadata['title'] = core_props.title
        metadata['author'] = core_props.author
        metadata['subject'] = core_props.subject
        metadata['created'] = str(core_props.created) if core_props.created else None
        metadata['modified'] = str(core_props.modified) if core_props.modified else None

        return metadata
